import os
import csv
import logging
import hashlib
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from collections import defaultdict
from pathlib import Path
from enum import Enum
from openpyxl import load_workbook
from PIL import Image
import io
from docx import Document
from docx.shared import Inches
import tempfile
from ttkthemes import ThemedTk

# 配置参数
DEFAULT_THRESHOLD = 5
HASH_SIZE = 8

class HashAlgorithm(Enum):
    PHASH = "感知哈希 (pHash)"
    DHASH = "差异哈希 (dHash)"
    AHASH = "平均哈希 (aHash)"

class ImageComparatorApp(ThemedTk):
    def __init__(self):
        super().__init__(theme="arc")
        self.title("Excel图片查重专业版")
        self.geometry("800x600")
        self.configure_layout()
        self.setup_logging()
        self.running = False
        self.pause_flag = threading.Event()

    def configure_layout(self):
        """创建现代化界面布局"""
        main_frame = ttk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 输入控制区
        input_frame = ttk.LabelFrame(main_frame, text="文件设置")
        input_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

        ttk.Label(input_frame, text="Excel文件夹:").grid(row=0, column=0, sticky="w")
        self.input_folder = ttk.Entry(input_frame, width=40)
        self.input_folder.grid(row=0, column=1, padx=5)
        ttk.Button(input_frame, text="浏览", command=self.select_input_folder).grid(row=0, column=2)

        ttk.Label(input_frame, text="输出报告:").grid(row=1, column=0, sticky="w")
        self.output_file = ttk.Entry(input_frame, width=40)
        self.output_file.grid(row=1, column=1, padx=5)
        ttk.Button(input_frame, text="选择", command=self.select_output_file).grid(row=1, column=2)

        # 算法设置区
        settings_frame = ttk.LabelFrame(main_frame, text="算法设置")
        settings_frame.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)

        ttk.Label(settings_frame, text="哈希算法:").grid(row=0, column=0, sticky="w")
        self.algorithm = ttk.Combobox(settings_frame, values=[a.value for a in HashAlgorithm])
        self.algorithm.current(0)
        self.algorithm.grid(row=0, column=1)

        ttk.Label(settings_frame, text="重复阈值:").grid(row=1, column=0, sticky="w")
        self.threshold = ttk.Spinbox(settings_frame, from_=2, to=100, width=5)
        self.threshold.set(DEFAULT_THRESHOLD)
        self.threshold.grid(row=1, column=1)

        # 输出选项
        self.word_var = tk.BooleanVar(value=True)
        self.csv_var = tk.BooleanVar()
        ttk.Checkbutton(settings_frame, text="生成Word报告", variable=self.word_var).grid(row=2, column=0, sticky="w")
        ttk.Checkbutton(settings_frame, text="生成CSV报告", variable=self.csv_var).grid(row=2, column=1, sticky="w")

        # 进度条
        self.progress = ttk.Progressbar(main_frame, orient=tk.HORIZONTAL, mode='determinate')
        self.progress.grid(row=1, column=0, columnspan=2, sticky="ew", pady=10)

        # 日志显示
        log_frame = ttk.LabelFrame(main_frame, text="处理日志")
        log_frame.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=5)
        self.log_area = ScrolledText(log_frame, height=10)
        self.log_area.pack(fill=tk.BOTH, expand=True)

        # 控制按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=3, column=0, columnspan=2, pady=10)
        self.start_btn = ttk.Button(btn_frame, text="开始处理", command=self.start_processing)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.pause_btn = ttk.Button(btn_frame, text="暂停", command=self.toggle_pause, state=tk.DISABLED)
        self.pause_btn.pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="退出", command=self.destroy).pack(side=tk.RIGHT, padx=5)

        # 布局权重配置
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=1)

    def setup_logging(self):
        """配置日志系统"""
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        
        # 文件日志
        file_handler = logging.FileHandler("image_checker.log")
        file_handler.setFormatter(formatter)
        
        # 界面日志
        text_handler = logging.StreamHandler(self.LogStream(self.log_area))
        text_handler.setFormatter(formatter)
        
        self.logger = logging.getLogger()
        self.logger.addHandler(file_handler)
        self.logger.addHandler(text_handler)
        self.logger.setLevel(logging.INFO)

    class LogStream:
        """自定义日志输出流"""
        def __init__(self, text_widget):
            self.text_widget = text_widget
            
        def write(self, message):
            self.text_widget.insert(tk.END, message)
            self.text_widget.see(tk.END)
            
        def flush(self):
            pass

    def select_input_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.input_folder.delete(0, tk.END)
            self.input_folder.insert(0, folder)

    def select_output_file(self):
        filetypes = [("Word文档", "*.docx"), ("CSV文件", "*.csv")]
        filename = filedialog.asksaveasfilename(filetypes=filetypes, defaultextension=".docx")
        if filename:
            self.output_file.delete(0, tk.END)
            self.output_file.insert(0, filename)

    def start_processing(self):
        if self.running:
            return
            
        self.running = True
        self.pause_btn.config(state=tk.NORMAL)
        self.start_btn.config(state=tk.DISABLED)
        
        threading.Thread(target=self.process_files, daemon=True).start()

    def toggle_pause(self):
        if self.pause_flag.is_set():
            self.pause_flag.clear()
            self.pause_btn.config(text="暂停")
        else:
            self.pause_flag.set()
            self.pause_btn.config(text="继续")

    def process_files(self):
        try:
            excel_files = self.find_excel_files()
            total_files = len(excel_files)
            
            with tempfile.TemporaryDirectory() as temp_dir:
                image_db = defaultdict(list)
                hasher = ImageHasher(self.get_selected_algorithm())
                
                for idx, path in enumerate(excel_files, 1):
                    if self.pause_flag.is_set():
                        while self.pause_flag.is_set():
                            threading.Event().wait(0.5)
                    
                    self.process_single_file(path, temp_dir, image_db, hasher)
                    self.update_progress(idx / total_files * 100)
                
                self.generate_reports(image_db, temp_dir)
                messagebox.showinfo("完成", "处理完成！")
        
        except Exception as e:
            self.logger.error(f"处理失败: {str(e)}")
            messagebox.showerror("错误", str(e))
        finally:
            self.running = False
            self.start_btn.config(state=tk.NORMAL)
            self.pause_btn.config(state=tk.DISABLED)
            self.progress["value"] = 0

    def find_excel_files(self):
        folder = Path(self.input_folder.get())
        return [p for p in folder.rglob("*") if p.suffix.lower() in (".xlsx", ".xlsm")]

    def process_single_file(self, path, temp_dir, image_db, hasher):
        try:
            self.logger.info(f"正在处理: {path.name}")
            wb = load_workbook(path)
            
            for sheet in wb.worksheets:
                for image in sheet._images:
                    img = Image.open(io.BytesIO(image._data()))
                    img_hash = hasher.calculate_hash(img)
                    image_db[img_hash].append(path)
                    
                    temp_path = Path(temp_dir) / f"{img_hash}.png"
                    if not temp_path.exists():
                        img.save(temp_path)
        except Exception as e:
            self.logger.error(f"处理失败 {path}: {str(e)}")

    def generate_reports(self, image_db, temp_dir):
        output_path = Path(self.output_file.get())
        threshold = int(self.threshold.get())
        
        if self.word_var.get():
            self.generate_word_report(image_db, output_path, temp_dir, threshold)
        if self.csv_var.get():
            self.generate_csv_report(image_db, output_path.with_suffix('.csv'), threshold)

    def generate_word_report(self, image_db, output_path, temp_dir, threshold):
        doc = Document()
        doc.add_heading('图片查重报告', 0)
        
        duplicates = sorted(
            [(h, files) for h, files in image_db.items() if len(files) >= threshold],
            key=lambda x: len(x[1]), 
            reverse=True
        )
        
        for h, files in duplicates:
            img_path = Path(temp_dir) / f"{h}.png"
            if img_path.exists():
                doc.add_heading(f'重复图片 (出现次数: {len(files)})', 1)
                doc.add_picture(str(img_path), width=Inches(3.5))
                doc.add_paragraph("出现位置：")
                for f in files:
                    doc.add_paragraph(f"• {f.name}")
        
        doc.save(output_path)

    def generate_csv_report(self, image_db, output_path, threshold):
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['哈希值', '出现次数', '文件列表'])
            for h, files in image_db.items():
                if len(files) >= threshold:
                    writer.writerow([h, len(files), '; '.join(str(p) for p in files)])

    def get_selected_algorithm(self):
        return [a for a in HashAlgorithm if a.value == self.algorithm.get()][0]

    def update_progress(self, value):
        self.progress["value"] = value
        self.update_idletasks()

class ImageHasher:
    def __init__(self, algorithm):
        self.algorithm = algorithm
        
    def calculate_hash(self, img):
        img = img.convert("L").resize((HASH_SIZE, HASH_SIZE), Image.LANCZOS)
        
        if self.algorithm == HashAlgorithm.PHASH:
            return self.phash(img)
        elif self.algorithm == HashAlgorithm.DHASH:
            return self.dhash(img)
        elif self.algorithm == HashAlgorithm.AHASH:
            return self.ahash(img)

    def phash(self, img):
        pixels = list(img.getdata())
        avg = sum(pixels) / len(pixels)
        bits = "".join(['1' if p > avg else '0' for p in pixels])
        return hashlib.md5(bits.encode()).hexdigest()

    def dhash(self, img):
        diff = []
        for row in range(HASH_SIZE):
            for col in range(HASH_SIZE-1):
                diff.append(img.getpixel((col, row)) > img.getpixel((col+1, row)))
        return hashlib.md5(''.join(str(int(b)) for b in diff).encode()).hexdigest()

    def ahash(self, img):
        avg = sum(img.getdata()) / (HASH_SIZE**2)
        bits = "".join(['1' if p > avg else '0' for p in img.getdata()])
        return hashlib.md5(bits.encode()).hexdigest()

if __name__ == "__main__":
    app = ImageComparatorApp()
    app.mainloop()